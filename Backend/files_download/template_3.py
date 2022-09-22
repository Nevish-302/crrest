from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from para import add_para


def template_3():
    # Declaring the document variable
    document = Document()

    # Defining Sections of the document
    sections = document.sections

    # Setting margins
    for section in sections:
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Table
    table = document.add_table(3, 3)
    table.autofit = False
    table.rows[2].height = Inches(8.2)
    table.cell(1, 0).merge(table.cell(1, 1)).merge(table.cell(1, 2))
    table.cell(2, 1).width = Inches(0.3)
    table.cell(0, 1).width = Inches(0.3)
    
    #To remove the additional line in the cells
    for cell in table._cells:
        cell._element.clear_content()
    
    #Form Data
    dictList = {
        "About": "Diligent Computer Science student (8.5 CGPA) from JSSATE, Noida with proven development and communication skills. Seeking a Software Developer position.",
        "Education": [
            {
                "year": "2014 - 2018",
                "degree": "B.tech Computer Science \n(Data Science)",
                "college": "JSSATEN",
                "location": "NOIDA",
            },
            {
                "year": "2020 - 2021",
                "degree": "+2",
                "college": "Navkis Residential Pre-University College",
                "location": "Bangalore, Karnataka",
            },
            {
                "year": "2018 - 2019",
                "degree": "10th",
                "college": "DAV PUBLIC SCHOOL",
                "location": "VASANT KUNJ NEW DELHI",
            },
        ],
        "Personal_Skills": [
            "Communication",
            "Teamwork",
            "Creativity",
            "Leadership",
            "Management",
        ],
        "Professional_Skills": [
            "Photoshop",
            "Python, Pandas",
            "Web Dev"
            "MongoDB",
            "Javascript",
        ],
        "Experience": [
            {
                "position": "Intern",
                "company": "HCL Tech. ",
                "year": "2024 - 2025",
                "Experience": "Learnt teamwork along with many skills where I contributed to several projects on the team with experienced professionals. Learnt comradery and got a more real experience to life.",
            },
        ],
        "Contact": {
            "Email": "ssatvik125@gmail.com",
            "Phone": "6366615165",
            "Website": "7wik.com",
            "Linkedin": "Satvik-8762016",
        },
    }
    
    #Adding Name
    add_para(
        "SATVIK",
        cell=table.cell(1, 0),
        document=document,
        font_size=60,
        line_spacing=0.6,
        font_name="Calibri Light",
        color=RGBColor(26, 23, 24),
        style_name="heading",
        font_bold=True,
    )
    
    #To increase the height of the mentioned cell 
    add_para(
        " ",
        cell=table.cell(0, 0),
        document=document,
        style_name="space",
    )
    
    #Preferred Job
    add_para(
        "SOFTWARE DEVELOPER",
        cell=table.cell(1, 0),
        document=document,
        font_size=28,
        line_spacing=0.6,
        font_name="Calibri Light",
        color=RGBColor(54, 48, 50),
        style_name="subheading",
        font_bold=True,
    )
    
    # Profile Pic
    pic = add_para("", style_name="picture", cell=table.cell(2, 0))
    run = pic.add_run()
    run.add_picture(
        r"C:\Users\Nevish302\Desktop\dv-modified.png",
        width=Inches(2.2),
    )
    
    # About
    add_para(
        "ABOUT ME",
        cell=table.cell(2, 0),
        font_size=22,
        font_name="Calibri Light",
        color=RGBColor(200, 200, 200),
        document=document,
        style_name="sun",
        font_bold=True,
    )
    
    #About Content
    add_para(
        dictList["About"],
        cell=table.cell(2, 0),
        font_name="Calibri Light",
        document=document,
        font_size=12,
        color=RGBColor(100, 100, 100),
        style_name="moon",
        font_bold=True,
    )
    
    # Cell Divider. Tbh this is the best option i found. If anyone has a better one, change it.
    add_para(
        "_______________",
        cell=table.cell(2, 0),
        font_name="Calibri Light",
        document=document,
        font_size=22,
        color=RGBColor(100, 100, 100),
        style_name="___",
        font_bold=True,
    )
    
    # Contact
    add_para(
        "CONTACT ME",
        cell=table.cell(2, 0),
        font_name="Calibri Light",
        document=document,
        font_size=22,
        color=RGBColor(200, 200, 200),
        style_name="contact",
        font_bold=True,
    )
    # So that the style name keeps changing
    i = 0
    # Contact content
    for key in dictList["Contact"]:
        add_para(
            f"\n{key}",
            cell=table.cell(2, 0),
            font_name="Calibri Light",
            font_size=12,
            document=document,
            color=RGBColor(150, 150, 150),
            style_name=str(i),
            font_bold=True,
        )
        i += 1
        add_para(
            f'{dictList["Contact"][key]}',
            cell=table.cell(2, 0),
            font_name="Calibri Light",
            font_size=12,
            document=document,
            color=RGBColor(100, 100, 100),
            style_name=str(i),
            font_bold=False,
        )
        i += 1
    #Newline
    add_para(
        "",
        cell=table.cell(2, 0),
        font_name="Calibri Light",
        font_size=12,
        document=document,
        color=RGBColor(100, 100, 100),
        style_name="oreally",
        font_bold=False,
    )
    
    # Education
    add_para(
        "EDUCATION",
        cell=table.cell(2, 2),
        font_name="Calibri Light",
        document=document,
        font_size=22,
        color=RGBColor(20, 20, 20),
        style_name="Education",
        font_bold=True,
    )
    
    # Education table
    tab = table.cell(2, 2).add_table(len(dictList["Education"]), 3)
    
    # Clearing the already existing content in the cell 
    for cell in tab._cells:
        cell._element.clear_content()
    
    # To add data in different cells
    j = 0
    
    # Education Content
    for degree in dictList["Education"]:
        tab.cell(j, 0)._element.clear_content()
        tab.cell(j, 0).width = Inches(1.21)
        tab.cell(j, 1).width = Inches(0.18)
        tab.cell(j, 2).width = Inches(3.23)

        add_para(
            degree["year"],
            cell=tab.cell(j, 0),
            font_name="Calibri Light",
            font_size=12,
            font_bold=True,
            document=document,
            color=RGBColor(100, 100, 100),
            style_name=str(i),
        )
        i += 1
        add_para(
            f"{degree['degree']}",
            cell=tab.cell(j, 2),
            font_name="Calibri Light",
            font_size=14,
            font_bold=True,
            document=document,
            color=RGBColor(100, 100, 100),
            style_name=str(i),
        )
        i += 1
        add_para(
            f"{degree['college']} ,{degree['location']}",
            cell=tab.cell(j, 2),
            font_name="Calibri Light",
            font_size=12,
            font_bold=False,
            document=document,
            color=RGBColor(100, 100, 100),
            style_name=str(i),
        )
        j += 1
        i += 1
    
    # Cell Divider. Tbh this is the best option i found. If anyone has a better one, change it.
    add_para(
        "___________________________________________",
        cell=table.cell(2, 2),
        font_name="Calibri Light",
        document=document,
        font_size=15,
        color=RGBColor(20, 20, 20),
        style_name="345",
        font_bold=True,
    )
    
    # Experience
    add_para(
        "EXPERIENCE",
        cell=table.cell(2, 2),
        font_name="Calibri Light",
        document=document,
        font_size=22,
        color=RGBColor(20, 20, 20),
        style_name="Experience",
        font_bold=True,
    )
    
    #Experience table
    tab = table.cell(2, 2).add_table(len(dictList["Experience"]), 3)
    
    # to clear already existing content
    for cell in tab._cells:
        cell._element.clear_content()
        
    # To add the content in different rows
    j = 0
    for degree in dictList["Experience"]:
        tab.cell(j, 0).width = Inches(1.21)
        tab.cell(j, 1).width = Inches(0.18)
        tab.cell(j, 2).width = Inches(3.23)

        add_para(
            degree["year"],
            cell=tab.cell(j, 0),
            font_name="Calibri Light",
            font_size=12,
            font_bold=True,
            document=document,
            color=RGBColor(100, 100, 100),
            style_name=str(i),
        )
        i += 1
        add_para(
            f"{degree['company']}",
            cell=tab.cell(j, 2),
            font_name="Calibri Light",
            font_size=12,
            font_bold=False,
            document=document,
            color=RGBColor(100, 100, 100),
            style_name=str(i),
        )
        i += 1
        add_para(
            f"{degree['position']}\n{degree['Experience']}",
            cell=tab.cell(j, 2),
            font_name="Calibri Light",
            font_size=12,
            font_bold=True,
            document=document,
            color=RGBColor(100, 100, 100),
            style_name=str(i),
        )
        j += 1
        i += 1
    
    # Cell Divider. Tbh this is the best option i found. If anyone has a better one, change it.
    add_para(
        "___________________________________________",
        cell=table.cell(2, 2),
        font_name="Calibri Light",
        document=document,
        font_size=15,
        color=RGBColor(20, 20, 20),
        style_name="123",
        font_bold=True,
    )
    
    # Setting the column width in the table
    tab = table.cell(2, 2).add_table(1, 3)
    tab.cell(0, 0).width = Inches(2.22)
    tab.cell(0, 1).width = Inches(0.18)
    tab.cell(0, 2).width = Inches(2.22)
    for cell in tab._cells:
        cell._element.clear_content()
    
    # Adding personal skills
    add_para(
        "PERSONAL SKILLS",
        cell=tab.cell(0, 0),
        font_name="Calibri Light",
        document=document,
        font_size=16,
        color=RGBColor(20, 20, 20),
        style_name=f"{i}",
        font_bold=True,
    )
    i += 1
    
    # Adding professional skills
    add_para(
        "PROFESSIONAL SKILLS",
        cell=tab.cell(0, 2),
        font_name="Calibri Light",
        document=document,
        font_size=16,
        color=RGBColor(20, 20, 20),
        style_name=f"{i}",
        font_bold=True,
    )
    i += 1
    for skill in dictList["Personal_Skills"]:

        add_para(
            f"{skill}",
            cell=tab.cell(0, 0),
            font_name="Calibri Light",
            font_size=12,
            font_bold=False,
            document=document,
            color=RGBColor(100, 100, 100),
            style_name=str(i),
        )
        i += 1
    for skill in dictList["Professional_Skills"]:

        add_para(
            f"{skill}",
            cell=tab.cell(0, 2),
            font_name="Calibri Light",
            font_size=12,
            font_bold=False,
            document=document,
            color=RGBColor(100, 100, 100),
            style_name=str(i),
        )
        i += 1

    # Coloring Cells
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="262324"/>'.format(nsdecls("w")))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="262324"/>'.format(nsdecls("w")))
    table.rows[2].cells[0]._tc.get_or_add_tcPr().append(shading_elm_2)

    # Saving the document
    document.save("template_3.docx")


template_3()
